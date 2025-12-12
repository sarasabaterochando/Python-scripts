import pandas as pd
from docx import Document
from pathlib import Path
from pypdf import PdfReader, PdfWriter
from reportlab.pdfgen import canvas
from reportlab.lib.pagesizes import letter
from io import BytesIO
import os
import logging
import traceback
import shutil
import streamlit as st

# === Load Excel data ===
data = pd.read_excel(file_xlsx)

# === Folder containing the Word templates ===
docs_source_path = path_doc_files

# === Base PDF template ===
pdf_template = Path(path_pdf_file)

# Output folder
output_folder = Path(path_final)

# Docx keywords
docx_name = "{{NAME}}"
docx_id_number = "{{ID_NUMBER}}"
docx_role = "{{ROLE}}"
docx_date = "{{DATE}}"

# Log file path
log_path = Path(path_log)

# === Logging configuration ===
logging.basicConfig(
    filename=log_path / 'transformation_log.log',
    level=logging.ERROR,
    format='%(asctime)s - %(levelname)s - %(message)s'
)


def log_function(func):
    """Decorator to catch and log errors."""
    def wrapper(*args, **kwargs):
        try:
            return func(*args, **kwargs)
        except Exception as e:
            error_info = traceback.format_exc()
            logging.error(f"Error in {func.__name__} with args {args}, kwargs {kwargs}:\n{error_info}")
            st.error(f"An error has been logged in {func.__name__}: {e}")
            return None
    return wrapper


@log_function
def save_docx(output_name, doc):
    st.write('-' * 20)
    st.write(f'\nSaving file: {output_name}')
    output_path = output_folder / f"{output_name}.docx"
    output_path.parent.mkdir(exist_ok=True)
    doc.save(output_path)
    st.write(f'Successfully saved file {output_path}')


@log_function
def process_word_files(name, id_number, role, date, document_type):
    word_path = Path(f'{docs_source_path}{document_type}')
    file_number = 0

    # === Iterate over .docx files ===
    for docx_file in word_path.glob("*.docx"):
        file_number += 1
        doc = Document(docx_file)

        # === Replace placeholders inside paragraphs ===
        for p in doc.paragraphs:
            if "{{NAME}}" in p.text:
                p.text = p.text.replace("{{NAME}}", str(name))
            if "{{ID_NUMBER}}" in p.text:
                p.text = p.text.replace("{{ID_NUMBER}}", str(id_number))
            if "{{ROLE}}" in p.text:
                p.text = p.text.replace("{{ROLE}}", str(role))
            if "{{DATE}}" in p.text:
                p.text = p.text.replace("{{DATE}}", str(date))

        # === Replace placeholders inside tables ===
        for table in doc.tables:
            for row in table.rows:
                for cell in row.cells:
                    if "{{NAME}}" in cell.text:
                        cell.text = cell.text.replace("{{NAME}}", str(name))
                    if "{{ID_NUMBER}}" in cell.text:
                        cell.text = cell.text.replace("{{ID_NUMBER}}", str(id_number))
                    if "{{ROLE}}" in cell.text:
                        cell.text = cell.text.replace("{{ROLE}}", str(role))
                    if "{{DATE}}" in cell.text:
                        cell.text = cell.text.replace("{{DATE}}", str(date))

        output_name = f'{name}_{id_number}_{file_number}'
        save_docx(output_name, doc)


@log_function
def fill_pdf_files(name, id_number, date):
    """
    Fill PDF files using the provided data.
    """
    reader = PdfReader(pdf_template)
    page = reader.pages[0]

    # Create a temporary PDF overlay with text
    packet = BytesIO()
    can = canvas.Canvas(packet, pagesize=letter)

    # Coordinates must be adjusted manually
    can.setFont("Helvetica", 10)
    can.drawString(100, 705, name)
    can.drawString(50, 705, id_number)

    can.drawString(65, 65, name)
    can.drawString(100, 112, date)
    can.drawString(50, 112, "City")  # Placeholder city

    can.save()
    packet.seek(0)

    # Merge overlay with original PDF
    overlay_reader = PdfReader(packet)
    overlay_page = overlay_reader.pages[0]
    page.merge_page(overlay_page)

    # Save final PDF
    writer = PdfWriter()
    writer.add_page(page)

    output_path = output_folder / f"{name}_{id_number}.pdf"
    with open(output_path, "wb") as f:
        writer.write(f)


@log_function
def process_excel(document_type, file_type):
    # === Loop through Excel rows ===
    for _, row in data.iterrows():
        name = row["Name"]
        id_number = row["IdNumber"]
        role = row["Role"]
        date = row["Date"]

        if file_type == "word":
            process_word_files(name, id_number, role, date, document_type)
        elif file_type == "pdf":
            fill_pdf_files(name, id_number, date)


@log_function
def clear_output_folder():
    # Delete files and folders inside output directory
    for item in output_folder.iterdir():
        st.write(f'Trying to delete: {item}')
        try:
            if item.is_file():
                os.remove(item)
            elif item.is_dir():
                shutil.rmtree(item)
        except Exception as e:
            st.error(f"❌ Error deleting {item}: {e}")


def render_ui():
    st.title("Document Processing Application")
    st.header("Automatically fill DOCX and PDF files using an Excel template")

    st.write("*" * 30)
    st.markdown(
        "<b style='color:navy'>Press <i>Run</i> to start the process.</b>",
        unsafe_allow_html=True
    )


if __name__ == '__main__':

    render_ui()

    document_type = st.radio("Document type to fill:", ["type_1", "type_2", "type_3"], index=0)

    file_type = st.radio("File type to generate:", ["pdf", "word"], index=0)

    try:
        run_button = st.button('Run')
        if run_button:
            clear_output_folder()
            process_excel(document_type, file_type)
            st.success('Process completed successfully')
    except Exception as e:
        st.error(f'Error executing process: {e}')
        traceback.print_exc()


